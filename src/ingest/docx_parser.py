"""DOCX parser module."""

from pathlib import Path
from typing import Dict

from src.common.logger import get_logger

# Try to import python-docx
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DOCXParser:
    """Parse DOCX files and extract text."""
    
    def __init__(self):
        self.logger = get_logger(__name__)
        if not HAS_DOCX:
            self.logger.warning(
                "python-docx not available. Install with: pip install python-docx"
            )
    
    def parse(self, file_path: str) -> Dict:
        """
        Parse DOCX file and extract text.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Dictionary with text and metadata
        """
        docx_path = Path(file_path)
        
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )
        
        try:
            text = self._extract_text(str(docx_path))
            metadata = self._extract_metadata(str(docx_path))
            
            if not text or len(text.strip()) < 10:
                self.logger.warning(
                    f"Extracted text from DOCX is very short or empty: {file_path}"
                )
            
            return {
                "text": text,
                "metadata": {
                    "file_path": str(docx_path),
                    "file_size": docx_path.stat().st_size,
                    **metadata
                }
            }
        
        except Exception as e:
            self.logger.error(f"Failed to parse DOCX {file_path}: {e}")
            # Return empty text instead of raising error
            return {
                "text": "",
                "metadata": {
                    "file_path": str(docx_path),
                    "error": str(e)
                }
            }
    
    def _extract_text(self, docx_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(docx_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        
        except Exception as e:
            self.logger.warning(f"Error reading DOCX {docx_path}: {e}")
            return ""
    
    def _extract_metadata(self, docx_path: str) -> Dict:
        """Extract metadata from DOCX file."""
        metadata = {}
        
        try:
            doc = Document(docx_path)
            core_props = doc.core_properties
            
            if core_props:
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                })
        
        except Exception as e:
            self.logger.debug(f"Could not extract metadata: {e}")
        
        return metadata

